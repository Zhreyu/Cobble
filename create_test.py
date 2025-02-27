import os
import random
import string
from fpdf import FPDF
from docx import Document
import openpyxl


# Function to generate random text
def generate_random_text(length=100):
    """Generates random text of a given length."""
    return ''.join(random.choices(string.ascii_letters + string.digits + string.punctuation + ' ', k=length))


# Content categories for each file
content_categories = {
    'Physics': [
        "Newton's laws of motion and their applications",
        "Theory of relativity by Einstein",
        "Electromagnetism and the study of fields",
        "Quantum mechanics and wave-particle duality"
    ],
    'Chemistry': [
        "Periodic table and the properties of elements",
        "Acids, Bases, and pH levels",
        "Chemical reactions and stoichiometry",
        "Organic chemistry and hydrocarbons"
    ],
    'Mathematics': [
        "Calculus and derivatives",
        "Probability theory and statistics",
        "Linear algebra and matrices",
        "Trigonometry and angle calculations"
    ],
    'General Ideas': [
        "Brainstorming ideas for a new project",
        "Steps to start a new business",
        "Creative writing prompts and exercises",
        "Team building activities and techniques"
    ],
    'To Do List': [
        "Complete homework assignments",
        "Buy groceries for the week",
        "Call the dentist for an appointment",
        "Plan the weekend trip with friends"
    ]
}


# Function to select random category content
def get_random_category_content(category):
    """Selects random content from a specific category."""
    return random.choice(content_categories.get(category, ["Random content"]))


# Function to create .txt file
def create_txt_file(file_path, category):
    """Create a sample .txt file with specific category content."""
    content = f"Category: {category}\n\n{get_random_category_content(category)}"
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)


# Function to create .md file
def create_md_file(file_path, category):
    """Create a sample .md (Markdown) file with specific category content."""
    content = f"# {category} Sample\n\n{get_random_category_content(category)}"
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)


# Function to create .pdf file
def create_pdf_file(file_path, category):
    """Create a sample .pdf file with specific category content."""
    content = f"Category: {category}\n\n{get_random_category_content(category)}"
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, content)
    pdf.output(file_path)


# Function to create .docx file
def create_docx_file(file_path, category):
    """Create a sample .docx file with specific category content."""
    content = f"Category: {category}\n\n{get_random_category_content(category)}"
    doc = Document()
    doc.add_heading(f"Sample {category} Content", 0)
    doc.add_paragraph(content)
    doc.save(file_path)


# Function to create .xls file
def create_xls_file(file_path, category):
    """Create a sample .xls file with specific category content."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"{category} Sheet"
    ws['A1'] = f"Category: {category}"
    ws['A2'] = get_random_category_content(category)
    wb.save(file_path)


# Function to create a number of sample files
def create_sample_files(base_path, num_files=5):
    """Create the specified number of sample files in the given base directory."""
    if not os.path.exists(base_path):
        os.makedirs(base_path)

    # List of possible categories to assign to files
    categories = list(content_categories.keys())

    for i in range(num_files):
        category = random.choice(categories)
        file_name = f"sample_{i+1}"
        
        create_txt_file(os.path.join(base_path, f"{file_name}.txt"), category)
        create_md_file(os.path.join(base_path, f"{file_name}.md"), category)
        create_pdf_file(os.path.join(base_path, f"{file_name}.pdf"), category)
        create_docx_file(os.path.join(base_path, f"{file_name}.docx"), category)
        create_xls_file(os.path.join(base_path, f"{file_name}.xlsx"), category)

    print(f"{num_files} sample files created in the '{base_path}' directory.")


if __name__ == "__main__":
    base_path = "test"
    num_files = 3  # Number of files to create
    create_sample_files(base_path, num_files)
